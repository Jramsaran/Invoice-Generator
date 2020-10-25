# finding a file's directory address on Mac https://stackoverflow.com/questions/3324486/finding-a-files-directory-address-on-a-mac
import pdfplumber
import PySimpleGUI as sg
from constants import *
import traceback
import pandas as pd
from docx import Document
from docx2pdf import convert
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
import os

def search_patterns_JPR(text_file):
       
    for index, line in enumerate(text_file.split("\n")):
        
        line2 = line
        line3 = line
        
        line = re_JPR.search(line)
        line3 = rit_extra_kosten_JPR.search(line3)
        
        if line:
        
            # print(line.group(7))
                    
            JPR_data[0].append(index)
            
            JPR_data[1].append(line.group(3))
                
            JPR_data[2].append(line.group(1))
                    
            JPR_data[3].append(line.group(5))
            
            JPR_data[4].append("€ " + line.group(7))
            
            JPR_data[5].append(line.group(7))
            
        #extra kosten en aantal en mss nog meer
 
        if line3:
            
            JPR_extra_data[0].append(index)
            
            JPR_extra_data[1].append("€ " + line3.group(1))
            
            JPR_extra_data[2].append(line3.group(2))

        
        if charterkosten_JPR.match(line2):
            
            search_patterns_JPR.charterkost_euro = line2.split()[-1]
            
            search_patterns_JPR.charterkost = float(line2.split()[-1].replace(".", "").replace(",", "."))

            
        elif extrakosten_JPR.match(line2):
            
            search_patterns_JPR.extrakost_euro = line2.split()[-1]
            
            search_patterns_JPR.extrakost = float(line2.split()[-1].replace(".", "").replace(",", "."))
            
            
        elif totaalbedragen_JPR.match(line2):
            
            search_patterns_JPR.totaal_bedrag_euro = line2.split()[-1]
            
            search_patterns_JPR.totaal_bedrag = float(line2.split()[-1].replace(".", "").replace(",", "."))
              
def search_patterns_JSR(text_file):
       
    # Search patterns from regular expressions in text
     
    for line in text_file.split("\n"):
                
        if datums_JSR.match(line):
                    
            JSR_data[0].append(line.split()[1])
                
        elif ritnummers_JSR.match(line):
                    
            JSR_data[1].append(line.split()[-1])
                
        elif containernummers_JSR.match(line):
                    
            JSR_data[2].append(line.split()[0] + " " + line.split()[1])
                    
        elif rit_bedragen_JSR.match(line):
                    
            JSR_data[3].append("€ " + line.split()[0])
            
            JSR_data[4].append(line.split()[0])
            
        elif totaalbedrag_JSR.match(line):
            
            search_patterns_JSR.totaal_bedrag_euro = line.split()[-1]
            
            search_patterns_JSR.totaal_bedrag = float(line.split()[-1].replace(".", "").replace(",", "."))


def string_to_float(header, i):
    
    i = int(i)
    
    for bedrag in range(len(header[i])):
        
        header[i][bedrag] = float(header[i][bedrag].replace(".", "").replace(",", "."))
  
        
def test_amount(list_of_amounts, total):
    
    test_amount.eindbedrag = 0
    test_amount.check = 0
    
    for bedrag in range(len(list_of_amounts)):
        
        test_amount.eindbedrag += list_of_amounts[bedrag]
        
    
    if round(test_amount.eindbedrag, 2) == round(total, 2):
        
        test_amount.check = 1
        


def read_pdf(file_name):
    
    with pdfplumber.open(file_name) as pdf:

        compiled_text = ""        
        
        for n in range(len(pdf.pages)):
                    
            page = pdf.pages[n]
            text = page.extract_text()
            compiled_text += text
            
        read_pdf.ct = compiled_text

def match_index(first_list, second_list):
    #first list is extra data
     
    diesel = ""
    chassishuur = ""
    diesel_bedrag = ""
    chassishuur_bedrag = ""
    
    for index in range(len(first_list[0])):
        
        if any(word in "diesel" for word in first_list[2][index].split()):
            
            diesel = "Dieseltoeslag week " + first_list[2][index].split()[int(config["DEFAULT"]["aantal woorden na diesel"])]
            diesel_bedrag = first_list[1][index]

        elif any(word in "chassishuur" for word in first_list[2][index].split()):
            
            chassishuur = "Chassishuur"
            chassishuur_bedrag = first_list[1][index]

        else:
        
            for index2 in range(len(second_list[0])):
                
                if int(first_list[0][index]) - 1 == int(second_list[0][index2]):
                    
                    second_list[2][index2] = second_list[2][index2] + "\n" + first_list[2][index].capitalize()
                    second_list[4][index2] = second_list[4][index2] + "\n" + first_list[1][index]
    
    match_index.diesel_chassis = diesel + "\n" + chassishuur
    match_index.diesel_chassis_bedrag = diesel_bedrag + "\n" + chassishuur_bedrag

# def make_word():
    
#     doc = docx.Document(word_template)
                    
#     # add a table to the end and create a reference variable
#     # extra row is so we can add the header row
                    
#     t = doc.add_table(df.shape[0] + 1, df.shape[1])
    
#     # t12 = doc.add_table(1, 1)
    
#     t2 = doc.add_table(df2.shape[0], df2.shape[1])
    
#     if context == 1:
        
#         eindbedracht = search_patterns_JPR.totaal_bedrag_euro
        
#     elif context > 1:
            
#         eindbedracht = search_patterns_JSR.totaal_bedrag_euro
        
#         eindtekst = f"Wij verzoeken u vriendelijk het verschuldigde bedrag \
#             van € {eindbedracht} binnen 30 dagen over te \
#                 maken naar IBAN: {gegevens[context]['iban']} \
#                     ten name van {gegevens[context]['Bedrijfsnaam']} onder vermelding van \
#                         factuurnummer: {factuurnummer}"
                
#         eindtekst1 = "Wij verzoeken u vriendelijk het verschuldigde bedrag van"             
#         eindtekst2 = f" € {eindbedracht}"
#         eindtekst3 = " binnen 30 dagen over te maken naar IBAN: "
#         eindtekst4 = f"{gegevens[context]['iban']}"
#         eindtekst5 = f" ten name van {gegevens[context]['Bedrijfsnaam']} onder vermelding van "
#         eindtekst6 = f"factuurnummer: {factuurnummer}"
                    
                    
#         #haalt dubbele spaties weg
#         # eindtekst = " ".join(eindtekst.split())
                    
#         lege_regels = doc.add_paragraph("\n")
                    
#         eindbericht = doc.add_paragraph(f"{eindtekst1}")
#         eindbericht.add_run(f"{eindtekst2}")
#         eindbericht.add_run(f"{eindtekst3}")
#         eindbericht.add_run(f"{eindtekst4}")
#         eindbericht.add_run(f"{eindtekst5}")
#         eindbericht.add_run(f"{eindtekst6}")
        
#         eindbericht.runs[1].font.bold = True
#         eindbericht.runs[3].font.bold = True
#         eindbericht.runs[5].font.bold = True
        
#         # add the header rows for df
                  
#         for j in range(df.shape[-1]):
#             t.cell(0,j).text = df.columns[j]
                        
#         # add the rest of the data frame for df
                    
#         for i in range(df.shape[0]):
#             for j in range(df.shape[-1]):
#                 t.cell(i+1,j).text = str(df.values[i,j])
                            
#         # t.cell(-1,(df.shape[-1] - 2)).text = str(df2.values[0,0])
#         # t.cell(-1,df.shape[-1] - 1).text = str(df2.values[0,1])
#         print(df, df2)
#         # add data for df2
                    
#         for i in range(df2.shape[0]):
#             for j in range(df2.shape[-1]):
#                 t2.cell(i,j).text = str(df2.values[i,j])
                            
#                     # for cell in t_columns[0].cells:
#                     #     cell.width = Inches(0.5)

#         t.style = config["STYLE"]["TABLE"]
#         # t12.style = config["STYLE"]["TABLE2"]
#         t2.style = config["STYLE"]["TABLE2"]
#         # t.alignment = WD_TABLE_ALIGNMENT.CENTER
#         # t2.alignment = WD_TABLE_ALIGNMENT.CENTER

#         # DOCUMENT OPSLAAN
#         for i in range(df.shape[0] + 1):
                    
#             t.cell(i, (df.shape[-1] - 1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
#         for j in range(df2.shape[0]):
                        
#             t2.cell(j, (df2.shape[-1] - 1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
#         if int(config["STYLE"]["KOSTEN_LOCATIE"]) == 2:
                    
#             if context == 1:
                    
#                 t2.cell((df2.shape[0] - 1), 2).text = 'Totaal charterkosten\nTotaal extra kosten\nTotaal generaal'
                            
#                 t2.cell((df2.shape[0] - 1), 2).paragraphs[0].runs[0].font.bold = True
                            
#             else:
                            
#                 t2.cell((df2.shape[0] - 1), 2).text = 'Totaal charterkosten'
                            
#                 t2.cell((df2.shape[0] - 1), 2).paragraphs[0].runs[0].font.bold = True
                            
#         for k in range(df.shape[-1]):
                        
#             t2.cell((df2.shape[0] - 1), k).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                                        
#             doc.save(values["Opslaan als..."] + ".docx")

if __name__ == '__main__':
    
    while True:
        event, values = window.read()
        # print(event, values)
        
        factuurnummer = values["-FN-"]
        word_template = values["-wt-"]
        table_style = values["-ts-"]
        
        # Eerste drie knoppen
        
        if event in (sg.WIN_CLOSED, 'Exit'):
            
            break
        
        elif event == '-JPR-':                # if the normal button that changes color and text
            down1 = not down1
            window['-JPR-'].update(button_color= ('#000000', '#C2D4D8') if down1 else "white on brown")
            window['-FN-'].update("" if down1 else factuurnummer_dict[1])
            
            # Het bedrijf en de bijbehorende dictionary kiezen
            
            if down1 == False:
                
                context = 1
                
            else:
                
                context = 0
            
            if down2 == False or down3 == False:
                down2 = True
                window['-JSR-'].update(button_color= ('#000000', '#C2D4D8') if down2 else "white on brown")

                down3 = True
                window['-GJ-'].update(button_color= ('#000000', '#C2D4D8') if down3 else "white on brown")
            
        elif event == '-JSR-':                # if the normal button that changes color and text
            down2 = not down2
            window['-JSR-'].update(button_color= ('#000000', '#C2D4D8') if down2 else "white on brown")
            window['-FN-'].update("" if down2 else factuurnummer_dict[2])
            
            if down2 == False:
                
                context = 2
                
            else:
                
                context = 0
            
            if down1 == False or down3 == False:
                down1 = True
                window['-JPR-'].update(button_color= ('#000000', '#C2D4D8') if down1 else "white on brown")

                down3 = True
                window['-GJ-'].update(button_color= ('#000000', '#C2D4D8') if down3 else "white on brown")
            
        elif event == '-GJ-':                # if the normal button that changes color and text
            down3 = not down3
            window['-GJ-'].update(button_color= ('#000000', '#C2D4D8') if down3 else "white on brown")
            window['-FN-'].update("" if down3 else factuurnummer_dict[3])
            
            if down3 == False:
                
                context = 3
                
            else:
                
                context = 0
            
            if down1 == False or down2 == False:
                down1 = True
                window['-JPR-'].update(button_color= ('#000000', '#C2D4D8') if down1 else "white on brown")

                down2 = True
                window['-JSR-'].update(button_color= ('#000000', '#C2D4D8') if down2 else "white on brown")
        
        if event == "-MF-" and values["Opslaan als..."] != "":
            
            # Reading PDF file
        
            sg.user_settings_set_entry('-last WT-', values['-wt-'])
            window['-wt-'].update(values=list(set(sg.user_settings_get_entry('-WT-', []))))
            
            sg.user_settings_set_entry('-last TS-', values['-ts-'])
            window['-ts-'].update(values=list(set(sg.user_settings_get_entry('-TS-', []))))
            
            if context >= 1:
                
                try:
                
                    window['-STATUS-'].update("\nFactuur gemaakt.", text_color = "green")
                    
                    read_pdf(values["Kies bestand"])
                        
                    if context == 1:
                    
                        search_patterns_JPR(read_pdf.ct)
                        
                        string_to_float(JPR_data, 5)
                        
                        test_amount(JPR_data[5], search_patterns_JPR.charterkost)
                        
                        match_index(JPR_extra_data, JPR_data)
                        
                        tweede_tabel[0].append("\n")
                        tweede_tabel[0].append("\n")
                        
                        if int(config["STYLE"]["KOSTEN_LOCATIE"]) == 0:
                            tweede_tabel[int(config["STYLE"]["KOSTEN_LOCATIE"])].append("Totaal charterkosten" + "\n" \
                                                   + "Totaal extra kosten" + "\n" + "Totaal generaal")
                        
                        if int(config["STYLE"]["ANDERE_LOCATIE"]) == 0:
                            tweede_tabel[int(config["STYLE"]["ANDERE_LOCATIE"])].append("\n")
                        # tweede_tabel[2].append("Totaal extra kosten")
                        # tweede_tabel[2].append("Totaal generaal")
                        
                        tweede_tabel[3].append(match_index.diesel_chassis_bedrag)
                        tweede_tabel[3].append("\n")
                        tweede_tabel[3].append("€ " + search_patterns_JPR.charterkost_euro \
                                               + "\n" + "€ " + search_patterns_JPR.extrakost_euro \
                                                   + "\n" + "€ " + search_patterns_JPR.totaal_bedrag_euro)
                        # tweede_tabel[3].append("€ " + search_patterns_JPR.extrakost_euro)
                        # tweede_tabel[3].append("€ " + search_patterns_JPR.totaal_bedrag_euro)
                       
                        tweede_tabel[2].append(match_index.diesel_chassis)
                        tweede_tabel[2].append("\n")
                        
                        if int(config["STYLE"]["KOSTEN_LOCATIE"]) == 2:
                            
                            tweede_tabel[int(config["STYLE"]["KOSTEN_LOCATIE"])].append('')

                        if int(config["STYLE"]["ANDERE_LOCATIE"]) == 2:
                            tweede_tabel[int(config["STYLE"]["ANDERE_LOCATIE"])].append("\n")
                        
                        
                        for entry in range(len(tweede_tabel[3])):
                            
                            tweede_tabel[1].append("\n")
                        
                        df = pd.DataFrame(JPR_data[1:5], index = headers).transpose()
                        df2 = pd.DataFrame(tweede_tabel[::]).transpose()
                        
                    else:
                        
                        search_patterns_JSR(read_pdf.ct)
    
                        string_to_float(JSR_data, 4)
                            
                        test_amount(JSR_data[4], search_patterns_JSR.totaal_bedrag)
                        
                        tweede_tabel[0].append("\n")
                        tweede_tabel[1].append("\n")
                        tweede_tabel[2].append("Totaal charterkosten")
                        tweede_tabel[3].append("€ " + search_patterns_JSR.totaal_bedrag_euro)
                        
                        df = pd.DataFrame(JSR_data[:4], index = headers).transpose()        
                        df2 = pd.DataFrame(tweede_tabel[::]).transpose()
                        
                    # SCHRIJVEN IN WORD DOC
                    
                    doc = Document(word_template)
                    
                    # add a table to the end and create a reference variable
                    # extra row is so we can add the header row
                    
                    t = doc.add_table(df.shape[0] + 1, df.shape[1])
                    
                    # t12 = doc.add_table(1, 1)
                    
                    t2 = doc.add_table(df2.shape[0], df2.shape[1])
                    
                    if context == 1:
                    
                        eindbedracht = search_patterns_JPR.totaal_bedrag_euro
                        
                    elif context > 1:
                        
                        eindbedracht = search_patterns_JSR.totaal_bedrag_euro
                    
                    eindtekst = f"Wij verzoeken u vriendelijk het verschuldigde bedrag \
                            van € {eindbedracht} binnen 30 dagen over te \
                                maken naar IBAN: {gegevens[context]['iban']} \
                                ten name van {gegevens[context]['Bedrijfsnaam']} onder vermelding van \
                                    factuurnummer: {factuurnummer}"
                    
                    eindtekst1 = "Wij verzoeken u vriendelijk het verschuldigde bedrag van"             
                    eindtekst2 = f" € {eindbedracht}"
                    eindtekst3 = " binnen 30 dagen over te maken naar IBAN: "
                    eindtekst4 = f"{gegevens[context]['iban']}"
                    eindtekst5 = f" ten name van {gegevens[context]['Bedrijfsnaam']} onder vermelding van "
                    eindtekst6 = f"factuurnummer: {factuurnummer}"
                    
                    
                    #haalt dubbele spaties weg
                    # eindtekst = " ".join(eindtekst.split())
                    
                    lege_regels = doc.add_paragraph("\n\n\n\n\n\n")
                    
                    eindbericht = doc.add_paragraph(f"{eindtekst1}")
                    eindbericht.add_run(f"{eindtekst2}")
                    eindbericht.add_run(f"{eindtekst3}")
                    eindbericht.add_run(f"{eindtekst4}")
                    eindbericht.add_run(f"{eindtekst5}")
                    eindbericht.add_run(f"{eindtekst6}")
                    
                    eindbericht.runs[1].font.bold = True
                    eindbericht.runs[3].font.bold = True
                    eindbericht.runs[5].font.bold = True
                    
                    # add the header rows for df
                   
                    for j in range(df.shape[-1]):
                        t.cell(0,j).text = df.columns[j]
                        
                    # add the rest of the data frame for df
                    
                    for i in range(df.shape[0]):
                        for j in range(df.shape[-1]):
                            t.cell(i+1,j).text = str(df.values[i,j])
                            
                    # t.cell(-1,(df.shape[-1] - 2)).text = str(df2.values[0,0])
                    # t.cell(-1,df.shape[-1] - 1).text = str(df2.values[0,1])
                    # print(df, df2)
                    # add data for df2
                    
                    for i in range(df2.shape[0]):
                        for j in range(df2.shape[-1]):
                            t2.cell(i,j).text = str(df2.values[i,j])
                            
                    # for cell in t_columns[0].cells:
                    #     cell.width = Inches(0.5)

                    t.style = table_style
                    # t12.style = config["STYLE"]["TABLE2"]
                    t2.style = table_style
                    # t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # t2.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # DOCUMENT OPSLAAN
                    for i in range(df.shape[0] + 1):
                    
                        t.cell(i, (df.shape[-1] - 1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                    for j in range(df2.shape[0]):
                        
                        t2.cell(j, (df2.shape[-1] - 1)).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                    if int(config["STYLE"]["KOSTEN_LOCATIE"]) == 2:
                    
                        if context == 1:
                    
                            t2.cell((df2.shape[0] - 1), 2).text = 'Totaal charterkosten\nTotaal extra kosten\nTotaal generaal'
                            
                            t2.cell((df2.shape[0] - 1), 2).paragraphs[0].runs[0].font.bold = True
                            
                        else:
                            
                            t2.cell((df2.shape[0] - 1), 2).text = 'Totaal charterkosten'
                            
                            t2.cell((df2.shape[0] - 1), 2).paragraphs[0].runs[0].font.bold = True
                            
                    for k in range(df.shape[-1]):
                        
                        t2.cell((df2.shape[0] - 1), k).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                                        
                    # Update template
                    
                    #centering bottom text
                    eindbericht.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.save(word_template.split(".")[0] + "_temp" + ".docx")
                    
                    doc2 = DocxTemplate(word_template.split(".")[0] + "_temp" + ".docx")
                    doc2.render(gegevens[context])
                    
                    doc2.save(values["Opslaan als..."] + ".docx")
                    os.remove(word_template.split(".")[0] + "_temp" + ".docx")
                    
                    factuur_gemaakt = 1
                    
                except Exception:
                    
                    traceback.print_exc()
                    
                    window['-STATUS-'].update("\nError", text_color = "red")
                
            else:
                
                window['-STATUS-'].update("Selecteer het bedrijf.", text_color = "red")
             
        elif event == "-MF-" and values["Opslaan als..."] == "":
            
            window['-STATUS-'].update("\nKlik eerst op 'Opslaan als...'", text_color = "red")
            
        if event == "-PDF-" and values["Opslaan als..."] != "" and factuur_gemaakt != 0:
            
            convert(values["Opslaan als..."] + ".docx", values["Opslaan als..."] + ".pdf")
            
            window['-STATUS-'].update("\nOpgeslagen als PDF.", text_color = "green")
            
        elif event == "-PDF-" and (values["Opslaan als..."] == "" or factuur_gemaakt == 0):
            
            window['-STATUS-'].update("\nKlik eerst op 'Opslaan als...' en 'Maak factuur'", text_color = "red")
            
            # print(read_pdf.ct)
            # print("%.2f, %.2f" % (search_patterns_JSR.totaal_bedrag, test_amount.eindbedrag))
        # print(JPR_data, JSR_data, JPR_extra_data, tweede_tabel)
            
        # print(context)
        # file_path_CK = "/".join(values["Kies bestand"].split("/")[:-1])
        # print(values["-FN-"], values["-wt-"])
        # file_path_WT = "/".join(values["-wt-"].split("/")[:-1])
        # print(file_path_CK, file_path_WT)
        # print(factuurnummer)
        # print(sg.user_settings_filename())
    
    window.close()
